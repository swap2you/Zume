from pathlib import Path

from zume import ingest

RESUME = """Aarav Mehta
Senior Automation Engineer — 9.2 years
Skills: Java, Selenium 4, TestNG, Cucumber, REST Assured, Oracle SQL, Jenkins.
"""


def test_extract_text_txt(tmp_path: Path):
    path = tmp_path / "resume.txt"
    path.write_text(RESUME, encoding="utf-8")
    assert "Aarav Mehta" in ingest.extract_text(path)


def test_extract_text_docx(tmp_path: Path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Aarav Mehta")
    doc.add_paragraph("9 years of automation experience")
    path = tmp_path / "resume.docx"
    doc.save(str(path))
    text = ingest.extract_text(path)
    assert "Aarav Mehta" in text
    assert "9 years" in text


def test_parse_resume_fields():
    profile = ingest.parse_resume(RESUME)
    assert profile.name == "Aarav Mehta"
    assert profile.experience_years == 9.2


def test_parse_name_skips_title_lines():
    text = "Resume\nSenior SDET\nRohan N\n7 years"
    assert ingest.parse_name(text) == "Rohan N"


def test_parse_experience_absent():
    assert ingest.parse_experience_years("Java, Selenium") is None


def test_parse_schedule_text():
    details = ingest.parse_schedule_text(
        "Candidate: Aarav Mehta\nDate: 2026-07-20\nTime: 10:00 AM ET\n"
        "Duration: 90 minutes\nInterviewers: Hiring Panel\nIgnored line\n")
    assert details["date"] == "2026-07-20"
    assert details["time"] == "10:00 AM ET"
    assert details["duration"] == "90 minutes"
    assert details["interviewers"] == "Hiring Panel"
