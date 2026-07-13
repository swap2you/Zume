"""Phase 6 — exercise-library protection, rotation and material separation."""

from zume import interview as iv
from zume.config import load_exercise_library, load_hiring_standard, load_theme
from zume.exercises import ExerciseSelector, fingerprint, load_exercises
from zume.ingest import parse_resume
from zume.screening import screen_resume

CORE_AREAS = ["java", "selenium", "rest_assured", "sql_oracle", "framework_design", "debugging"]

RICH_RESUME = """Aarav Mehta
Senior Automation Engineer with 9.2 years of experience.
Skills: Java, Selenium 4, TestNG, Cucumber, REST Assured, Oracle SQL, Jenkins, Appium, BrowserStack.
Built and maintained a Java Selenium framework, implemented REST Assured API chaining,
validated payments in Oracle SQL and integrated Jenkins. Mentored two engineers.
"""


def test_each_core_area_has_three_active_variants(repo_root):
    exercises = load_exercises(load_exercise_library(repo_root))
    for area in CORE_AREAS:
        active = [e for e in exercises[area] if e.status == "active"]
        assert len(active) >= 3, f"{area} needs >=3 active variants, has {len(active)}"


def test_every_exercise_has_required_fields(repo_root):
    exercises = load_exercises(load_exercise_library(repo_root))
    for area_exercises in exercises.values():
        for ex in area_exercises:
            assert ex.task
            assert ex.requirement_change_follow_up
            assert ex.debugging_follow_up
            assert ex.expected_reasoning
            assert ex.reference_solution
            assert ex.scoring_rubric
            assert ex.red_flags
            assert ex.independence_questions


def test_draft_and_retired_exercises_are_never_selected(repo_root):
    exercises = load_exercises(load_exercise_library(repo_root))
    selector = ExerciseSelector(exercises)
    picked = selector.pick("java", 3)
    ids = {e.id for e in picked}
    assert "java_stream_grouping" not in ids  # draft
    assert all(e.status == "active" for e in picked)


def test_rotation_prefers_unused_then_least_recently_used(repo_root):
    exercises = load_exercises(load_exercise_library(repo_root))
    first = ExerciseSelector(exercises).pick("selenium", 1)[0]
    # Mark the first choice as heavily used; the selector should rotate away.
    usage = {first.id: {"count": 5, "last_used": "2026-01-01T00:00:00+00:00"}}
    second = ExerciseSelector(exercises, usage=usage).pick("selenium", 1)[0]
    assert second.id != first.id


def test_candidate_history_is_avoided(repo_root):
    exercises = load_exercises(load_exercise_library(repo_root))
    first = ExerciseSelector(exercises).pick("sql_oracle", 1)[0]
    selector = ExerciseSelector(exercises, candidate_history={first.id})
    picked = selector.pick("sql_oracle", 1)[0]
    assert picked.id != first.id


def test_no_duplicate_variant_family_in_one_area(repo_root):
    exercises = load_exercises(load_exercise_library(repo_root))
    picked = ExerciseSelector(exercises).pick("rest_assured", 2)
    families = [e.variant_family for e in picked]
    assert len(families) == len(set(families))


def test_fingerprint_is_stable_and_task_specific(repo_root):
    assert fingerprint("Count duplicates ") == fingerprint("count   Duplicates")
    assert fingerprint("task a") != fingerprint("task b")


def test_candidate_sheet_excludes_reference_solutions(repo_root, tmp_path):
    theme = load_theme(repo_root)
    result = screen_resume(parse_resume(RICH_RESUME), load_hiring_standard(repo_root))
    library = load_exercise_library(repo_root)
    kit = iv.build_kit(library, result)

    candidate_path = tmp_path / "candidate.docx"
    interviewer_path = tmp_path / "interviewer.docx"
    iv.generate_candidate_exercise_sheet(theme, kit, candidate_path)
    iv.generate_interview_guide(theme, kit, library, interviewer_path)

    from docx import Document

    candidate_text = "\n".join(p.text for p in Document(str(candidate_path)).paragraphs)
    interviewer_text = "\n".join(p.text for p in Document(str(interviewer_path)).paragraphs)

    # A known reference-solution fragment must appear only in the interviewer copy.
    solution_fragments = [ex.reference_solution.splitlines()[0]
                          for ex in kit.exercises if ex.reference_solution]
    for fragment in solution_fragments:
        assert fragment not in candidate_text
    assert "Reference solution" not in candidate_text
    assert "Reference solution" in interviewer_text
