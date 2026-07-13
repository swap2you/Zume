from pathlib import Path

from docx import Document

from zume.documents import ZumeDocument
from zume.validation import validate_docx


def _build(theme, path: Path) -> None:
    doc = ZumeDocument(theme, "ATS Screening Report", "Candidate: Test Person")
    doc.banner("Proceed with technical screening.", kind="success", label="DECISION")
    doc.heading("Decision", 1)
    doc.key_values([("Decision", "Proceed"), ("Resume evidence coverage", "88%")])
    doc.heading("Resume evidence coverage — mandatory skills", 1)
    doc.table(["Skill", "Evidence"], [["Java", "Project"], ["Selenium", "Project"]])
    doc.heading("Risks and inconsistencies", 1)
    doc.bullets(["Verify REST Assured depth live."])
    doc.save(path)


def test_generated_docx_passes_structural_validation(theme, tmp_path: Path):
    path = tmp_path / "ATS_Screening.docx"
    _build(theme, path)
    report = validate_docx(path)
    assert report.ok, report.errors
    assert any("heading hierarchy" in p for p in report.passed)
    assert any("header" in p for p in report.passed)
    assert any("footer" in p for p in report.passed)
    assert any("page numbers" in p for p in report.passed)


def test_docx_contains_header_footer_and_page_field(theme, tmp_path: Path):
    path = tmp_path / "doc.docx"
    _build(theme, path)
    doc = Document(str(path))
    section = doc.sections[0]
    assert "ZUME" in section.header.paragraphs[0].text.upper()
    footer_xml = section.footer.paragraphs[0]._p.xml
    assert "Private hiring operations material" in section.footer.paragraphs[0].text
    assert "PAGE" in footer_xml


def test_plain_document_fails_validation(tmp_path: Path):
    doc = Document()
    doc.add_paragraph("Just text, no structure")
    path = tmp_path / "plain.docx"
    doc.save(str(path))
    report = validate_docx(path)
    assert not report.ok
    assert any("no headings" in e for e in report.errors)
    assert any("header missing" in e for e in report.errors)


def test_unfilled_placeholder_is_flagged(theme, tmp_path: Path):
    doc = ZumeDocument(theme, "Recruiter Feedback", "")
    doc.heading("Draft", 1)
    doc.paragraph("Hi Team, reviewed [Candidate] profile.")
    path = tmp_path / "draft.docx"
    doc.save(path)
    report = validate_docx(path)
    assert any("placeholder" in w for w in report.warnings)


def test_save_is_versioned(theme, tmp_path: Path):
    path = tmp_path / "report.docx"
    doc = ZumeDocument(theme, "Interview Schedule", "")
    doc.heading("Details", 1)
    doc.paragraph("First version")
    doc.heading("Interviewer preparation checklist", 1)
    doc.bullets(["Check link"])
    assert doc.save(path) is True
    doc2 = ZumeDocument(theme, "Interview Schedule", "")
    doc2.heading("Details", 1)
    doc2.paragraph("Second version with changes")
    doc2.heading("Interviewer preparation checklist", 1)
    doc2.bullets(["Check link"])
    assert doc2.save(path) is True
    assert path.with_name("report__v1.docx").exists()
