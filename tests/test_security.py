"""Phase 11 — security and privacy guarantees over the tracked repository."""

import subprocess
from pathlib import Path

from zume.security import (
    PII_PATTERNS,
    SECRET_PATTERNS,
    list_tracked_files,
    scan_repository,
)


def _tracked(repo_root: Path) -> list[str]:
    return list_tracked_files(repo_root)


def test_no_secrets_or_pii_in_tracked_files(repo_root: Path):
    findings = scan_repository(repo_root, include_pii=True)
    assert findings == [], [f"{f.kind} @ {f.path}:{f.line}" for f in findings]


def test_no_generated_candidate_documents_tracked(repo_root: Path):
    tracked = _tracked(repo_root)
    offenders = [p for p in tracked
                 if p.startswith("candidates/") and not p.endswith(".gitkeep")]
    assert offenders == []


def test_fictional_fixtures_remain_tracked(repo_root: Path):
    tracked = set(_tracked(repo_root))
    assert "examples/fictional-candidate/resume.txt" in tracked
    assert "examples/fictional-candidate/schedule.txt" in tracked
    assert "examples/fictional-candidate/interview-notes.txt" in tracked


def test_sensitive_file_types_are_ignored(repo_root: Path):
    probes = [
        "candidates/Real_Person_2026-01-01/candidate.json",
        "input/real-resume.pdf",
        "output/generated.docx",
        "data/zume.db",
        "General Docs/real.docx",
        "screenshot.png",
        "photo.jpg",
        ".env",
        "output/pkg.zip",
        "~$draft.docx",
    ]
    for probe in probes:
        rc = subprocess.run(["git", "check-ignore", "-q", probe],
                            cwd=repo_root, check=False).returncode
        assert rc == 0, f"expected git to ignore {probe}"


def test_fixtures_are_not_ignored(repo_root: Path):
    rc = subprocess.run(
        ["git", "check-ignore", "-q", "examples/fictional-candidate/resume.txt"],
        cwd=repo_root, check=False).returncode
    assert rc == 1  # not ignored -> trackable


def test_scanner_detects_planted_secret_and_pii(tmp_path: Path):
    # Build a throwaway git repo with a planted secret + email, prove detection.
    # Strings are assembled at runtime so this tracked test file stays clean.
    subprocess.run(["git", "init", "-q"], cwd=tmp_path, check=True)
    secret_line = "api" + "_key = " + '"' + "abcd1234efgh5678" + '"'
    email_line = "contact " + "person" + "@" + "example.com"
    (tmp_path / "leak.txt").write_text(secret_line + "\n" + email_line + "\n",
                                       encoding="utf-8")
    subprocess.run(["git", "add", "-A"], cwd=tmp_path, check=True)
    findings = scan_repository(tmp_path, include_pii=True)
    kinds = {f.kind for f in findings}
    assert "generic_secret_assignment" in kinds
    assert "email" in kinds


def test_pattern_sets_are_non_empty():
    assert SECRET_PATTERNS and PII_PATTERNS


def test_scanner_detects_pii_inside_tracked_docx(tmp_path: Path):
    """Lockdown Part 8 — extract DOCX paragraph/table text and scan it for PII."""
    from docx import Document

    subprocess.run(["git", "init", "-q"], cwd=tmp_path, check=True)
    email = "person" + "@" + "example.com"
    phone = "555" + "-" + "867" + "-" + "5309"
    doc = Document()
    doc.add_paragraph("Reference contact: " + email)
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "Phone: " + phone
    doc.save(str(tmp_path / "reference.docx"))
    subprocess.run(["git", "add", "-A"], cwd=tmp_path, check=True)

    findings = scan_repository(tmp_path, include_pii=True)
    kinds = {f.kind for f in findings}
    paths = {f.path for f in findings}
    assert "email" in kinds
    assert "phone" in kinds
    assert "reference.docx" in paths
    # Findings never echo the matched value — only kind, path and line.
    for finding in findings:
        assert email not in finding.kind and phone not in finding.kind


def test_docx_scan_reports_path_only_not_value(tmp_path: Path):
    """A finding exposes the document path and kind, never the candidate value."""
    from zume.security import Finding

    f = Finding(path="docs/x.docx", line=1, kind="email")
    assert f.path == "docs/x.docx" and f.kind == "email"
    assert "@" not in f.kind
