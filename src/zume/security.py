"""Lightweight repository secret and PII scanning over git-tracked files.

This scans only files that Git tracks, and only text files, so it never reads
ignored candidate material. It reports the *kind* and location of a finding, and
deliberately does not echo the matched candidate text.
"""

from __future__ import annotations

import re
import subprocess
from dataclasses import dataclass
from pathlib import Path

TEXT_SUFFIXES = {
    ".py", ".md", ".txt", ".yaml", ".yml", ".toml", ".cfg", ".ini", ".json",
    ".gitignore", ".env", "", ".rst",
}

# Files that legitimately discuss these topics (docs about privacy/secrets).
# They are still scanned for real secret *values*, only the keyword heuristics
# below are what these would trip, so we keep the high-signal patterns strict.

SECRET_PATTERNS: dict[str, re.Pattern[str]] = {
    "aws_access_key_id": re.compile(r"\bAKIA[0-9A-Z]{16}\b"),
    "private_key_block": re.compile(r"-----BEGIN (?:RSA |OPENSSH |EC )?PRIVATE KEY-----"),
    "generic_secret_assignment": re.compile(
        r"(?i)\b(?:api[_-]?key|secret|token|password|passwd|pwd)\b\s*[:=]\s*"
        r"['\"][^'\"\s]{8,}['\"]"),
    "bearer_token": re.compile(r"(?i)\bbearer\s+[A-Za-z0-9._-]{20,}\b"),
    "slack_token": re.compile(r"\bxox[baprs]-[0-9A-Za-z-]{10,}\b"),
}

PII_PATTERNS: dict[str, re.Pattern[str]] = {
    "email": re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"),
    "phone": re.compile(r"(?<!\d)(?:\+\d{1,3}[\s-]?)?\(?\d{3}\)?[\s.-]\d{3}[\s.-]\d{4}(?!\d)"),
}


@dataclass
class Finding:
    path: str
    line: int
    kind: str


def list_tracked_files(root: Path) -> list[str]:
    result = subprocess.run(
        ["git", "ls-files"], cwd=root, capture_output=True, text=True, check=False)
    if result.returncode != 0:
        return []
    return [line for line in result.stdout.splitlines() if line.strip()]


def _is_text(path: Path) -> bool:
    return path.suffix.lower() in TEXT_SUFFIXES or path.name == ".gitignore"


def _extract_docx_text(path: Path) -> str:
    """Extract paragraph and table text from a DOCX using python-docx.

    Returns an empty string if the file cannot be opened as a Word document."""
    try:
        from docx import Document
    except Exception:  # noqa: BLE001 - python-docx is a hard dependency, but stay safe
        return ""
    try:
        document = Document(str(path))
    except Exception:  # noqa: BLE001 - not a valid docx / unreadable
        return ""
    parts: list[str] = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _scan_text(text: str, patterns: dict[str, re.Pattern[str]], rel: str,
               findings: list[Finding]) -> None:
    for lineno, line in enumerate(text.splitlines(), start=1):
        for kind, pattern in patterns.items():
            if pattern.search(line):
                findings.append(Finding(path=rel, line=lineno, kind=kind))


def scan_repository(root: Path, include_pii: bool = True) -> list[Finding]:
    """Return secret/PII findings across git-tracked text and DOCX files.

    Only git-tracked files are inspected, so ignored candidate/private material
    is never read. Findings report the finding *kind* and document path only —
    never the matched value."""
    findings: list[Finding] = []
    patterns = dict(SECRET_PATTERNS)
    if include_pii:
        patterns.update(PII_PATTERNS)
    for rel in list_tracked_files(root):
        path = root / rel
        if not path.exists():
            continue
        if path.suffix.lower() == ".docx":
            text = _extract_docx_text(path)
            if text:
                _scan_text(text, patterns, rel, findings)
            continue
        if not _is_text(path):
            continue
        try:
            text = path.read_text(encoding="utf-8", errors="strict")
        except (UnicodeDecodeError, OSError):
            continue
        _scan_text(text, patterns, rel, findings)
    return findings
