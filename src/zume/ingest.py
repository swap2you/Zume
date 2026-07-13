"""Input ingestion: PDF, DOCX, TXT, pasted text and screenshot metadata."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from PIL import Image
from pypdf import PdfReader


@dataclass
class ExperienceAnalysis:
    """Deterministic analysis of overall-experience claims in a resume."""

    state: str  # "passed" | "failed" | "unknown" (resolved against the minimum later)
    years: float | None
    claims: list[float]
    detail: str


@dataclass
class ResumeProfile:
    name: str
    experience_years: float | None
    text: str
    skills_mentioned: list[str] = field(default_factory=list)
    experience: ExperienceAnalysis | None = None


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if suffix == ".docx":
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    if suffix in {".txt", ".md", ""}:
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"Unsupported resume format: {suffix}")


_YEARS_PATTERNS = [
    re.compile(r"(\d+(?:\.\d+)?)\s*\+?\s*(?:years|yrs)", re.IGNORECASE),
    re.compile(r"(?:experience|exp)\D{0,10}(\d+(?:\.\d+)?)", re.IGNORECASE),
]

_NAME_STOPWORDS = {
    "resume", "curriculum", "vitae", "cv", "profile", "senior", "junior",
    "engineer", "developer", "automation", "sdet", "qa", "lead", "analyst",
}


def parse_name(text: str) -> str:
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        words = re.findall(r"[A-Za-z][A-Za-z'.-]*", line)
        if not words:
            continue
        if any(w.lower() in _NAME_STOPWORDS for w in words):
            continue
        if 1 <= len(words) <= 4:
            return " ".join(words)
        break
    raise ValueError("Could not confidently detect the candidate name; provide --name.")


def parse_experience_years(text: str) -> float | None:
    """Best-effort single value (kept for backward compatibility)."""
    analysis = analyze_experience(text)
    return analysis.years


# Patterns that tie a year figure explicitly to *total* experience.
_TOTAL_EXPERIENCE_PATTERNS = [
    re.compile(
        r"(\d+(?:\.\d+)?)\s*\+?\s*(?:years|yrs)\.?\s+(?:of\s+)?"
        r"(?:overall\s+|total\s+|professional\s+|relevant\s+|industry\s+|combined\s+)?"
        r"(?:experience|exp\b)",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:overall|total|professional|relevant)?\s*experience\s*[:\-]?\s*"
        r"(?:of\s+)?(\d+(?:\.\d+)?)\s*\+?\s*(?:years|yrs)",
        re.IGNORECASE,
    ),
]

# A summary/title line such as "Senior Automation Engineer — 9.2 years".
_SUMMARY_ROLE_PATTERN = re.compile(
    r"(?:engineer|sdet|developer|tester|analyst|lead|architect|consultant)"
    r"[^\n]*?(\d+(?:\.\d+)?)\s*\+?\s*(?:years|yrs)",
    re.IGNORECASE,
)

# Any year-denominated figure (used to detect ambiguity / tenure-only mentions).
_ANY_YEARS_PATTERN = re.compile(r"(\d+(?:\.\d+)?)\s*\+?\s*(?:years|yrs)\b", re.IGNORECASE)

# Date ranges like "2016 - 2020", "Jan 2018 – Present".
_DATE_RANGE_PATTERN = re.compile(
    r"(?:19|20)\d{2}\s*(?:-|–|—|to)\s*(?:(?:19|20)\d{2}|present|current|now)",
    re.IGNORECASE,
)


def _plausible_years(value: float) -> bool:
    return 0 < value < 50


def analyze_experience(text: str) -> ExperienceAnalysis:
    """Classify overall-experience claims deterministically.

    Returns a state of "passed"/"failed"/"unknown" is deferred to the caller
    (which knows the configured minimum); here ``state`` is either "known" with a
    resolved ``years`` value, or "unknown" with a human-readable ``detail``.
    Never invents a number.
    """
    total_claims: list[float] = []
    for pattern in (*_TOTAL_EXPERIENCE_PATTERNS, _SUMMARY_ROLE_PATTERN):
        for match in pattern.finditer(text):
            value = float(match.group(1))
            if _plausible_years(value):
                total_claims.append(value)

    distinct = sorted({round(v, 1) for v in total_claims})

    if len(distinct) == 1:
        return ExperienceAnalysis(
            state="known", years=distinct[0], claims=distinct,
            detail=f"Single explicit total-experience claim: {distinct[0]:g} years.",
        )
    if len(distinct) > 1:
        spread = max(distinct) - min(distinct)
        if spread >= 1.0:
            return ExperienceAnalysis(
                state="unknown", years=None, claims=distinct,
                detail=("Conflicting total-experience claims found: "
                        + ", ".join(f"{d:g}" for d in distinct)
                        + " years. Manual confirmation required."),
            )
        # Values agree within a rounding margin; use the largest.
        return ExperienceAnalysis(
            state="known", years=max(distinct), claims=distinct,
            detail=f"Total-experience claims agree at ~{max(distinct):g} years.",
        )

    # No explicit total-experience claim. Distinguish ambiguous vs missing.
    tenure_mentions = [float(m.group(1)) for m in _ANY_YEARS_PATTERN.finditer(text)
                       if _plausible_years(float(m.group(1)))]
    date_ranges = _DATE_RANGE_PATTERN.findall(text)
    if date_ranges or tenure_mentions:
        return ExperienceAnalysis(
            state="unknown", years=None, claims=[],
            detail=("No stated total experience; only per-role tenure or date "
                    "ranges are present. Manual review required to total it up."),
        )
    return ExperienceAnalysis(
        state="unknown", years=None, claims=[],
        detail="No overall-experience statement found in the resume.",
    )


def parse_resume(text: str, name_override: str | None = None) -> ResumeProfile:
    name = name_override or parse_name(text)
    analysis = analyze_experience(text)
    return ResumeProfile(
        name=name,
        experience_years=analysis.years,
        text=text,
        experience=analysis,
    )


def image_metadata(path: Path) -> dict[str, str]:
    """Read reliable, non-OCR metadata from a screenshot."""
    meta: dict[str, str] = {"file": path.name}
    with Image.open(path) as img:
        meta["format"] = img.format or "unknown"
        meta["size"] = f"{img.width}x{img.height}"
        exif = img.getexif()
        if exif:
            # 306 = DateTime, 36867 = DateTimeOriginal
            for tag in (36867, 306):
                if tag in exif:
                    meta["captured_at"] = str(exif[tag])
                    break
    return meta


def parse_schedule_text(text: str) -> dict[str, str]:
    """Parse `Key: Value` schedule lines from pasted/confirmed text."""
    known = {
        "candidate": "candidate",
        "date": "date",
        "time": "time",
        "duration": "duration",
        "interviewers": "interviewers",
        "interviewer": "interviewers",
        "platform": "platform",
        "meeting": "platform",
        "meeting method": "platform",
        "timezone": "timezone",
        "time zone": "timezone",
        "tz": "timezone",
        "notes": "notes",
    }
    details: dict[str, str] = {}
    for line in text.splitlines():
        if ":" not in line:
            continue
        key, _, value = line.partition(":")
        mapped = known.get(key.strip().lower())
        if mapped and value.strip():
            details[mapped] = value.strip()
    return details
