"""Branded DOCX generation engine built on python-docx."""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table

from zume.candidate import versioned_write_bytes

BANNER_KINDS = {
    "info": ("secondary", "NOTE"),
    "success": ("success", "PROCEED"),
    "warning": ("warning", "ATTENTION"),
    "danger": ("danger", "RISK"),
}


def _hex_color(theme: dict[str, Any], key: str) -> str:
    return str(theme["brand"].get(key, "17365D"))


def _set_cell_shading(cell: Any, hex_fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _mark_header_row(row: Any) -> None:
    """Repeat this row as a table header on every page."""
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _prevent_row_split(row: Any) -> None:
    """Keep a table row's cells on one page when practical (Phase 13)."""
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    cant.set(qn("w:val"), "true")
    tr_pr.append(cant)


def _add_page_number_field(paragraph: Any) -> None:
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE \\* MERGEFORMAT")
    inner_run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    inner_run.append(text)
    fld.append(inner_run)
    paragraph._p.append(fld)


class ZumeDocument:
    """Builder that applies the Zume theme consistently."""

    def __init__(self, theme: dict[str, Any], title: str, subtitle: str = "") -> None:
        self.theme = theme
        self.doc: DocumentObject = Document()
        self._apply_base_styles()
        self._apply_header_footer()
        self._add_title(title, subtitle)

    # -- setup -------------------------------------------------------------

    def _apply_base_styles(self) -> None:
        layout = self.theme["layout"]
        style = self.doc.styles["Normal"]
        style.font.name = layout.get("font", "Aptos")
        style.font.size = Pt(float(layout.get("body_size_pt", 10.5)))
        style.paragraph_format.space_after = Pt(6)
        margin = Inches(float(layout.get("margins_inches", 0.75)))
        for section in self.doc.sections:
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin
        primary = RGBColor.from_string(_hex_color(self.theme, "primary"))
        secondary = RGBColor.from_string(_hex_color(self.theme, "secondary"))
        heading_font = layout.get("heading_font", "Aptos Display")
        for name, color in (("Heading 1", primary), ("Heading 2", secondary), ("Heading 3", secondary)):
            heading = self.doc.styles[name]
            heading.font.name = heading_font
            heading.font.color.rgb = color
            heading.paragraph_format.space_before = Pt(10)
            heading.paragraph_format.space_after = Pt(4)
            heading.paragraph_format.keep_with_next = True

    def _apply_header_footer(self) -> None:
        layout = self.theme["layout"]
        section = self.doc.sections[0]
        header_para = section.header.paragraphs[0]
        header_para.text = layout.get("header", "ZUME | AUTOMATION HIRING")
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header_para.runs:
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(_hex_color(self.theme, "primary"))
        footer_para = section.footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run(layout.get("footer", "Private hiring operations material") + "  |  ")
        run.font.size = Pt(8)
        _add_page_number_field(footer_para)

    def _add_title(self, title: str, subtitle: str) -> None:
        self.doc.add_heading(title, level=0)
        if subtitle:
            para = self.doc.add_paragraph(subtitle)
            for run in para.runs:
                run.font.color.rgb = RGBColor.from_string(_hex_color(self.theme, "accent"))
                run.font.size = Pt(9)

    # -- content -----------------------------------------------------------

    def heading(self, text: str, level: int = 1) -> None:
        self.doc.add_heading(text, level=level)

    def paragraph(self, text: str, bold: bool = False, size_pt: float | None = None) -> None:
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        if size_pt:
            run.font.size = Pt(size_pt)

    def bullets(self, items: list[str]) -> None:
        for item in items:
            self.doc.add_paragraph(item, style="List Bullet")

    def key_values(self, pairs: list[tuple[str, str]]) -> None:
        for key, value in pairs:
            para = self.doc.add_paragraph()
            run = para.add_run(f"{key}: ")
            run.bold = True
            para.add_run(value)

    def banner(self, text: str, kind: str = "info", label: str | None = None) -> None:
        """Color-coded banner that always pairs color with a text label."""
        color_key, default_label = BANNER_KINDS.get(kind, BANNER_KINDS["info"])
        fill = _hex_color(self.theme, color_key)
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        _set_cell_shading(cell, fill)
        para = cell.paragraphs[0]
        run = para.add_run(f"{label or default_label}: {text}")
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        run.font.size = Pt(10.5)
        self.doc.add_paragraph()

    def table(self, headers: list[str], rows: list[list[str]]) -> Table:
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        header_row = table.rows[0]
        _mark_header_row(header_row)
        header_fill = _hex_color(self.theme, "primary")
        for idx, text in enumerate(headers):
            cell = header_row.cells[idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(text)
            run.bold = True
            run.font.color.rgb = RGBColor.from_string("FFFFFF")
            _set_cell_shading(cell, header_fill)
        for row_idx, values in enumerate(rows):
            row = table.add_row()
            _prevent_row_split(row)
            if row_idx % 2 == 1:
                for cell in row.cells:
                    _set_cell_shading(cell, "F2F2F2")
            for idx, value in enumerate(values):
                row.cells[idx].text = str(value)
        self.doc.add_paragraph()
        return table

    def code_block(self, text: str, size_pt: float = 9.0) -> None:
        """Monospace, lightly shaded reference-code block that stays together."""
        for raw in text.splitlines() or [" "]:
            para = self.doc.add_paragraph()
            para.paragraph_format.keep_together = True
            para.paragraph_format.space_after = Pt(0)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), "F4F4F4")
            para._p.get_or_add_pPr().append(shd)
            run = para.add_run(raw or " ")
            run.font.name = "Consolas"
            run.font.size = Pt(size_pt)
        self.doc.add_paragraph()

    def page_break(self) -> None:
        self.doc.add_page_break()

    def spacer(self) -> None:
        self.doc.add_paragraph()

    # -- output ------------------------------------------------------------

    def to_bytes(self) -> bytes:
        buffer = io.BytesIO()
        self.doc.save(buffer)
        return buffer.getvalue()

    def save(self, path: Path, versioned: bool = True) -> bool:
        """Save the document.

        Legacy callers keep content-hash versioning; the v2 deliverables path
        passes ``versioned=False`` for an atomic replace with no ``__vN`` copies.
        """
        if versioned:
            return versioned_write_bytes(path, self.to_bytes())
        from zume.candidate import atomic_write_bytes

        atomic_write_bytes(path, self.to_bytes())
        return True
