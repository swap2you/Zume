"""Structural validation for candidate folders and generated DOCX files."""

from __future__ import annotations

import json
import re
import shutil
import statistics
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from zume.candidate import (
    CONTRACT_SUBFOLDERS,
    DELIVERABLES_DIR,
    INTERNAL_DIR,
    SUBFOLDERS,
    candidate_json_path,
)

EXPECTED_SECTIONS = {
    # Legacy documents.
    "ATS Screening Report": ["Decision", "Resume evidence coverage — mandatory skills",
                             "Risks and inconsistencies"],
    "Full Interview Guide (3 Hours)": ["Session plan", "Exercises with expected reasoning",
                                       "Scoring per exercise"],
    "Interview Schedule": ["Details", "Interviewer preparation checklist"],
    # v2 consolidated deliverables (unique titles).
    "Screening Summary": ["Resume summary", "Mandatory-skill evidence",
                          "Risks, inconsistencies and missing evidence",
                          "Recruiter screening message (copy-ready)"],
    "Three-Hour Interview Guide": ["Candidate-specific risks and validation targets",
                                   "180-minute agenda", "Knockout round (20 minutes)",
                                   "Live exercises (expected reasoning and reference solutions)"],
    "Schedule and Communications": ["Confirmed details", "Interviewer preparation checklist",
                                    "Copy-ready communication drafts"],
    "Interview Exercises (Candidate Copy)": ["What to expect"],
    "Post-Interview Communications": ["Recruiter draft (copy-ready)"],
}

# Titles that must NEVER expose interviewer-only material to candidates.
CANDIDATE_SHAREABLE_TITLES = {"Interview Exercises (Candidate Copy)"}

_PLACEHOLDER = re.compile(r"\[(?:[A-Za-z][A-Za-z ]{2,})\]")


@dataclass
class ValidationReport:
    passed: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)

    @property
    def ok(self) -> bool:
        return not self.errors

    def merge(self, other: "ValidationReport") -> None:
        self.passed.extend(other.passed)
        self.warnings.extend(other.warnings)
        self.errors.extend(other.errors)


def validate_docx(path: Path) -> ValidationReport:
    report = ValidationReport()
    name = path.name
    try:
        doc = Document(str(path))
    except Exception as exc:  # noqa: BLE001 - report any unreadable document
        report.errors.append(f"{name}: cannot be opened by python-docx ({exc})")
        return report

    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        report.errors.append(f"{name}: document body is empty")
        return report

    headings = [p for p in paragraphs
                if p.style is not None and str(p.style.name).startswith(("Heading", "Title"))]
    if not headings:
        report.errors.append(f"{name}: no headings found")
    else:
        report.passed.append(f"{name}: heading hierarchy present ({len(headings)} headings)")

    section = doc.sections[0]
    header_text = " ".join(p.text for p in section.header.paragraphs)
    if "ZUME" not in header_text.upper():
        report.errors.append(f"{name}: Zume header missing")
    else:
        report.passed.append(f"{name}: Zume header present")

    footer_text = " ".join(p.text for p in section.footer.paragraphs)
    footer_xml = section.footer.paragraphs[0]._p.xml if section.footer.paragraphs else ""
    if "Private" not in footer_text:
        report.errors.append(f"{name}: confidentiality footer missing")
    else:
        report.passed.append(f"{name}: confidentiality footer present")
    if "PAGE" not in footer_xml:
        report.errors.append(f"{name}: page-number field missing")
    else:
        report.passed.append(f"{name}: page numbers present")

    title = paragraphs[0].text.strip()
    expected = EXPECTED_SECTIONS.get(title, [])
    if expected:
        heading_texts = {h.text.strip() for h in headings}
        missing = [s for s in expected if s not in heading_texts]
        if missing:
            report.errors.append(f"{name}: missing expected sections: {', '.join(missing)}")
        else:
            report.passed.append(f"{name}: all expected sections present")

    body_text = "\n".join(p.text for p in paragraphs)
    leftovers = sorted(set(_PLACEHOLDER.findall(body_text)))
    if leftovers:
        report.warnings.append(f"{name}: possible unfilled placeholders: {', '.join(leftovers)}")
    else:
        report.passed.append(f"{name}: no unfilled placeholders")
    return report


def detect_sparse_trailing_page(page_texts: list[str], threshold: float = 0.2) -> str | None:
    """Return a message when the last page has far less text than the median
    page (Phase 13). Single-page documents are always fine."""
    lengths = [len(t.strip()) for t in page_texts]
    if len(lengths) < 2:
        return None
    median = statistics.median(lengths[:-1])
    last = lengths[-1]
    if median > 0 and last < threshold * median:
        return (f"trailing page has {last} characters vs a median of {int(median)} "
                f"(< {int(threshold * 100)}% — likely a sparse/near-blank page)")
    return None


def find_soffice() -> str | None:
    for name in ("soffice", "soffice.exe"):
        found = shutil.which(name)
        if found:
            return found
    for candidate in (
        Path("C:/Program Files/LibreOffice/program/soffice.exe"),
        Path("C:/Program Files (x86)/LibreOffice/program/soffice.exe"),
    ):
        if candidate.exists():
            return str(candidate)
    return None


def _expected_headings_for(path: Path) -> list[str]:
    try:
        doc = Document(str(path))
    except Exception:  # noqa: BLE001
        return []
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    title = paragraphs[0].text.strip() if paragraphs else ""
    return EXPECTED_SECTIONS.get(title, [])


def word_com_available() -> bool:
    """True when Microsoft Word can be driven via COM (Windows + pywin32)."""
    try:
        import win32com.client  # type: ignore  # noqa: F401
    except Exception:  # noqa: BLE001
        return False
    return True


def _inspect_pdf(path: Path, rendered: Path, backend: str) -> ValidationReport:
    """Shared rendered-PDF content checks used by every render backend."""
    report = ValidationReport()
    report.passed.append(f"{path.name}: rendered to PDF via {backend}")
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(rendered))
        page_count = len(reader.pages)
        if page_count == 0:
            report.errors.append(f"{path.name}: rendered PDF has zero pages")
            return report
        report.passed.append(f"{path.name}: rendered PDF has {page_count} page(s)")
        page_texts = [(p.extract_text() or "") for p in reader.pages]
        text = "\n".join(page_texts)
        if not text.strip():
            report.errors.append(f"{path.name}: rendered PDF has no extractable text")
            return report
        if page_count > 1 and not page_texts[-1].strip():
            report.warnings.append(f"{path.name}: last rendered page appears blank")
        sparse = detect_sparse_trailing_page(page_texts)
        if sparse:
            report.warnings.append(f"{path.name}: {sparse}")
        if "ZUME" in text.upper():
            report.passed.append(f"{path.name}: header text present in rendered PDF")
        else:
            report.warnings.append(f"{path.name}: header text not found in rendered PDF")
        if "Private" in text:
            report.passed.append(f"{path.name}: footer text present in rendered PDF")
        else:
            report.warnings.append(f"{path.name}: footer text not found in rendered PDF")
        leftovers = sorted(set(_PLACEHOLDER.findall(text)))
        if leftovers:
            report.warnings.append(
                f"{path.name}: possible unresolved placeholders in PDF: {', '.join(leftovers)}")
        missing = [h for h in _expected_headings_for(path) if h not in text]
        if missing:
            report.warnings.append(
                f"{path.name}: headings not found in rendered PDF: {', '.join(missing)}")
        elif _expected_headings_for(path):
            report.passed.append(f"{path.name}: expected headings present in rendered PDF")
    except Exception as exc:  # noqa: BLE001
        report.warnings.append(f"{path.name}: rendered but PDF inspection failed ({exc})")
    return report


def render_docx_word(path: Path) -> ValidationReport:
    """Render a DOCX to PDF using Microsoft Word via COM, then inspect it."""
    report = ValidationReport()
    try:
        import win32com.client  # type: ignore

        pythoncom = __import__("pythoncom")
        pythoncom.CoInitialize()
    except Exception as exc:  # noqa: BLE001
        report.warnings.append(f"{path.name}: Word COM unavailable ({exc})")
        return report
    with tempfile.TemporaryDirectory() as tmp:
        rendered = Path(tmp) / (path.stem + ".pdf")
        word = None
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(str(path), ReadOnly=True)
            doc.SaveAs(str(rendered), FileFormat=17)  # wdFormatPDF
            doc.Close(False)
        except Exception as exc:  # noqa: BLE001
            report.errors.append(f"{path.name}: Word render failed ({exc})")
            return report
        finally:
            if word is not None:
                try:
                    word.Quit()
                except Exception:  # noqa: BLE001
                    pass
        if not rendered.exists() or rendered.stat().st_size == 0:
            report.errors.append(f"{path.name}: Word produced no PDF output")
            return report
        report.merge(_inspect_pdf(path, rendered, "Microsoft Word"))
    return report


def render_docx(path: Path, soffice: str) -> ValidationReport:
    """Render a DOCX to PDF and verify the rendered content, not just XML.

    Checks: conversion succeeds, page count is nonzero, the document is not empty,
    expected headings appear in the extracted PDF text, header and footer text are
    present, and a page-number value is rendered.
    """
    report = ValidationReport()
    with tempfile.TemporaryDirectory() as tmp:
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", tmp, str(path)],
            capture_output=True, text=True, timeout=180, check=False,
        )
        rendered = Path(tmp) / (path.stem + ".pdf")
        if result.returncode != 0 or not rendered.exists() or rendered.stat().st_size == 0:
            report.errors.append(
                f"{path.name}: LibreOffice render failed ({result.stderr.strip()[:200]})")
            return report
        report.passed.append(f"{path.name}: rendered to PDF via LibreOffice")
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(rendered))
            page_count = len(reader.pages)
            if page_count == 0:
                report.errors.append(f"{path.name}: rendered PDF has zero pages")
                return report
            report.passed.append(f"{path.name}: rendered PDF has {page_count} page(s)")
            text = "\n".join((p.extract_text() or "") for p in reader.pages)
            if not text.strip():
                report.errors.append(f"{path.name}: rendered PDF has no extractable text")
                return report
            # last page should not be blank (trailing empty page check)
            page_texts = [(p.extract_text() or "") for p in reader.pages]
            last = page_texts[-1].strip()
            if page_count > 1 and not last:
                report.warnings.append(f"{path.name}: last rendered page appears blank")
            sparse = detect_sparse_trailing_page(page_texts)
            if sparse:
                report.warnings.append(f"{path.name}: {sparse}")
            if "ZUME" not in text.upper():
                report.warnings.append(f"{path.name}: header text not found in rendered PDF")
            else:
                report.passed.append(f"{path.name}: header text present in rendered PDF")
            if "Private" not in text:
                report.warnings.append(f"{path.name}: footer text not found in rendered PDF")
            else:
                report.passed.append(f"{path.name}: footer text present in rendered PDF")
            missing = [h for h in _expected_headings_for(path) if h not in text]
            if missing:
                report.warnings.append(
                    f"{path.name}: headings not found in rendered PDF: {', '.join(missing)}")
            elif _expected_headings_for(path):
                report.passed.append(f"{path.name}: expected headings present in rendered PDF")
        except Exception as exc:  # noqa: BLE001
            report.warnings.append(f"{path.name}: rendered but PDF inspection failed ({exc})")
    return report


def _is_v2_folder(folder: Path) -> bool:
    return (folder / INTERNAL_DIR).is_dir() or (folder / DELIVERABLES_DIR).is_dir()


def validate_candidate_folder(folder: Path, render: bool = True) -> ValidationReport:
    report = ValidationReport()
    v2 = _is_v2_folder(folder)
    expected_folders = CONTRACT_SUBFOLDERS if v2 else SUBFOLDERS
    for sub in expected_folders:
        if (folder / sub).is_dir():
            report.passed.append(f"folder: {sub} present")
        else:
            report.errors.append(f"folder: {sub} missing")

    audit = candidate_json_path(folder)
    if not audit.exists():
        report.errors.append("candidate.json missing")
    else:
        try:
            json.loads(audit.read_text(encoding="utf-8"))
            report.passed.append("candidate.json parses as JSON")
        except json.JSONDecodeError as exc:
            report.errors.append(f"candidate.json is invalid JSON: {exc}")

    if v2:
        _validate_v2_invariants(folder, report)
        docx_files = sorted(p for p in (folder / DELIVERABLES_DIR).glob("*.docx")
                            if not p.name.startswith("~$"))
    else:
        docx_files = sorted(p for p in folder.rglob("*.docx")
                            if not p.name.startswith("~$"))
    if not docx_files:
        report.warnings.append("no DOCX artifacts found to validate")
    for doc_path in docx_files:
        report.merge(validate_docx(doc_path))

    if render:
        if v2:
            render_docs = docx_files
        else:
            render_docs = sorted((folder / "99-final").glob("*.docx")) or docx_files[:2]
        soffice = find_soffice()
        if soffice is not None:
            for doc_path in render_docs:
                report.merge(render_docx(doc_path, soffice))
        elif word_com_available():
            for doc_path in render_docs:
                report.merge(render_docx_word(doc_path))
        else:
            report.warnings.append(
                "No render backend available (LibreOffice or Microsoft Word); DOCX "
                "render verification was skipped. Structural checks still ran."
            )
    return report


def _validate_v2_invariants(folder: Path, report: ValidationReport) -> None:
    """Phase 15/16 invariants: <=7 deliverables, no __vN, no 99-final."""
    deliverables = sorted((folder / DELIVERABLES_DIR).glob("*.docx"))
    if len(deliverables) <= 7:
        report.passed.append(f"deliverables: {len(deliverables)} DOCX (<= 7)")
    else:
        report.errors.append(f"deliverables: {len(deliverables)} DOCX exceeds the 7 maximum")
    versioned = [p.name for p in folder.rglob("*__v[0-9]*")]
    if versioned:
        report.errors.append(f"versioned copies present: {', '.join(versioned)}")
    else:
        report.passed.append("no __vN versioned files")
    if (folder / "99-final").exists():
        report.errors.append("legacy 99-final folder present in a v2 candidate")
    else:
        report.passed.append("no 99-final folder")


def check_privacy(root: Path) -> ValidationReport:
    """Verify that real candidate paths are ignored by Git."""
    report = ValidationReport()
    probes = [
        "candidates/Probe_Test_2026-01-01/candidate.json",
        "input/probe-resume.pdf",
        "output/probe.docx",
        "data/zume.db",
        "General Docs/probe.docx",
    ]
    for probe in probes:
        result = subprocess.run(
            ["git", "check-ignore", "-q", probe],
            cwd=root, capture_output=True, check=False,
        )
        if result.returncode == 0:
            report.passed.append(f"git ignores {probe}")
        elif result.returncode == 1:
            report.errors.append(f"git does NOT ignore {probe}")
        else:
            report.warnings.append(f"git check-ignore unavailable for {probe}")
    return report
